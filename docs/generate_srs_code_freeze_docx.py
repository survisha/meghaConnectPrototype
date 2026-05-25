from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(__file__).with_name(
    "SRS_Code_Freeze_IAS_Final_Approval_2026-05-25.docx"
)
FALLBACK_OUTPUT_PATH = Path(__file__).with_name(
    "SRS_Code_Freeze_AI_Transformation_IAS_Final_Approval_2026-05-25.docx"
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr_cells[idx], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr_cells[idx], "1F4E79")
        if widths:
            hdr_cells[idx].width = widths[idx]

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            if widths:
                cells[idx].width = widths[idx]
    document.add_paragraph()
    return table


def add_bullets(document, items, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        document.add_paragraph(item, style=style)


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def enable_update_fields_on_open(document):
    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_toc(document):
    document.add_heading("Table of Contents", level=1)
    paragraph = document.add_paragraph()
    field_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_begin)
    field_run._r.append(instr)
    field_run._r.append(fld_sep)
    paragraph.add_run("Table of Contents will update automatically when opened in Microsoft Word.")
    end_run = paragraph.add_run()
    end_run._r.append(fld_end)
    document.add_page_break()


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 22, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "365F91"),
        ("Heading 3", 11, "1F4E79"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "MeghaConnect - SRS and Code Freeze Document"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    add_page_number(section.footer.paragraphs[0])


def add_cover_page(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Software Requirements Specification (SRS)")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("1F4E79")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Code Freeze Document for Final Approval")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("365F91")

    project = document.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project.add_run("MeghaConnect - AI-enabled Citizen Engagement and Governance Workflow Platform")
    run.font.size = Pt(13)

    document.add_paragraph()
    add_table(
        document,
        ["Document Attribute", "Details"],
        [
            ("Prepared For", "IAS / Competent Authority - Final Approval"),
            ("Document Type", "SRS and Functional Code Freeze"),
            ("Version", "1.0"),
            ("Status", "Submitted for approval"),
            ("Date", "25 May 2026"),
            ("Prepared By", "MeghaConnect Project Team"),
        ],
        [Inches(2.2), Inches(4.8)],
    )

    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.add_run(
        "This document freezes the functional scope described herein. Any change after approval "
        "shall be handled through a formal change request, impact assessment, and revised approval."
    ).italic = True

    document.add_page_break()


def add_document_control(document):
    document.add_heading("1. Document Control", level=1)
    add_table(
        document,
        ["Version", "Date", "Description", "Author / Owner"],
        [
            (
                "1.0",
                "25 May 2026",
                "Initial code-freeze SRS prepared for final approval.",
                "MeghaConnect Project Team",
            )
        ],
        [Inches(1), Inches(1.2), Inches(3.5), Inches(1.8)],
    )

    document.add_heading("Approval Sign-off", level=2)
    add_table(
        document,
        ["Name", "Designation", "Signature", "Date", "Remarks"],
        [
            ("", "IAS / Competent Authority", "", "", ""),
            ("", "Project Owner", "", "", ""),
            ("", "Implementation Partner", "", "", ""),
        ],
        [Inches(1.4), Inches(1.7), Inches(1.2), Inches(1), Inches(2)],
    )


def add_intro(document):
    document.add_heading("2. Purpose and Scope", level=1)
    document.add_paragraph(
        "This SRS defines the frozen functional scope for the MeghaConnect system covering citizen "
        "registration, appointment management, scheme applications, staff dashboards, public "
        "identification, reporting, HCM actions, administration, and AI-assisted features."
    )
    document.add_paragraph(
        "The document is intended for final business approval prior to code freeze. It records the "
        "approved module flow, role responsibilities, dependencies, pending integrations, and key "
        "business rules."
    )

    document.add_heading("In Scope", level=2)
    add_bullets(
        document,
        [
            "Citizen registration with EPIC, Aadhaar, or No ID based registration paths.",
            "Citizen login, dashboard, appointment creation, scheme application, and gate pass download.",
            "DEO-assisted registration, walk-in appointment creation, and citizen detail updates.",
            "Public identification through search and future face-identification integration.",
            "Role-based staff dashboards, calendar scheduling, appointment review, reports, HCM actions, admin controls, and AI features.",
        ],
    )

    document.add_heading("Code Freeze Statement", level=2)
    add_bullets(
        document,
        [
            "All modules and flows listed in this document are considered frozen for approval.",
            "Items listed under Pending Items are accepted dependencies or post-template/post-approval tasks.",
            "New business rules or functional changes after approval shall be tracked as change requests.",
        ],
    )

    document.add_heading("Abbreviations", level=2)
    add_table(
        document,
        ["Term", "Meaning"],
        [
            ("AADHAAR", "Aadhaar based citizen identity verification"),
            ("AI", "Artificial Intelligence"),
            ("CMO", "Chief Minister Office user role"),
            ("DEO", "Data Entry Operator"),
            ("EPIC", "Electors Photo Identity Card"),
            ("HCM", "Hon'ble Chief Minister role / dashboard workflow"),
            ("IAS", "Indian Administrative Service / competent approving authority"),
            ("KYC", "Know Your Customer / identity verification"),
            ("OSD", "Officer on Special Duty"),
            ("OTP", "One Time Password"),
            ("OVSC", "Third-party Aadhaar QR authentication service"),
            ("OCR", "Optical Character Recognition / document reading"),
            ("QR", "Quick Response code used for entry pass validation"),
        ],
        [Inches(1.4), Inches(5.6)],
    )


def add_ai_transformation_vision(document):
    document.add_heading("3. AI Transformation Vision", level=1)
    document.add_paragraph(
        "MeghaConnect is positioned as an AI-enabled citizen engagement, appointment, scheme monitoring, "
        "and governance workflow platform for the Meghalaya CM Office. The platform is designed to support "
        "citizens, DEO users, CMO teams, Approvers, HCM users, OSD users, and administrators through a "
        "single integrated governance workflow."
    )
    document.add_paragraph(
        "AI capabilities are embedded within the existing MeghaConnect workflows rather than being limited "
        "to a separate standalone module. The objective is to reduce manual review effort, improve citizen "
        "guidance, support faster decision-making, and generate actionable insights for CMO, HCM, and "
        "Approver teams."
    )
    add_bullets(
        document,
        [
            "MeghaBot for citizen guidance during appointment and service journeys.",
            "AI Notes for reading uploaded documents and generating structured summaries.",
            "AI Dashboard Insights for trend visibility, scheme demand, district-wise patterns, and project category analysis.",
            "Face Identification for faster public identification, subject to future API and GPU infrastructure readiness.",
            "OCR / Document Understanding for extracting and interpreting uploaded scanned documents.",
            "Future predictive analytics and decision support for priority scoring, demand forecasting, and governance planning.",
        ],
    )


def add_current_workflow_challenges(document):
    document.add_heading("4. Current Workflow Challenges and AI Enablement", level=1)
    add_table(
        document,
        ["Current Challenge", "AI Enablement"],
        [
            ("Manual document review", "AI Notes and OCR summaries"),
            ("Citizen confusion during appointment creation", "MeghaBot assistance"),
            ("Manual search for citizen history", "AI-assisted public identification"),
            ("Large appointment volume", "AI dashboard and prioritization insights"),
            ("Manual reports", "AI-driven analytics and trends"),
            ("Future face-based identification need", "Face Recognition API"),
        ],
        [Inches(3.2), Inches(3.8)],
    )


def add_ai_architecture_overview(document):
    document.add_heading("5. AI Architecture Overview", level=1)
    document.add_paragraph(
        "The following logical flow shows how AI capabilities are integrated into MeghaConnect while retaining "
        "the platform as a formal SRS / Code Freeze scope for governance workflows."
    )
    add_table(
        document,
        ["Flow Layer", "Component / Function"],
        [
            ("1", "Citizen / Staff User"),
            ("2", "MeghaConnect Portal"),
            ("3", "AI Layer: MeghaBot, AI Notes Engine, OCR / Document Reader, Face Identification Engine, AI Dashboard Insights Engine"),
            ("4", "Business Services: registration, appointment, scheme, scheduling, approval, reporting, and administration services"),
            ("5", "Database: citizen records, appointments, schemes, documents, remarks, audit data, and reference tables"),
            ("6", "Reports and Dashboards: staff dashboard, HCM dashboard, scheme heatmap, analytics, audit trail, and AI insights"),
        ],
        [Inches(1.2), Inches(5.8)],
    )
    document.add_paragraph(
        "Textual flow: Citizen / Staff User -> MeghaConnect Portal -> AI Layer -> Business Services -> Database -> Reports and Dashboards."
    )


def add_ai_benefits_for_government(document):
    document.add_heading("6. AI Benefits for Government", level=1)
    add_table(
        document,
        ["Area", "Benefit"],
        [
            ("Citizen Service", "Faster guidance and reduced dependency on manual help"),
            ("CMO/Approver Review", "Faster document review and better prioritization"),
            ("HCM Dashboard", "Actionable decision support"),
            ("Reports", "Trend and demand visibility"),
            ("Public Identification", "Faster citizen lookup"),
            ("Governance", "Data-driven planning and monitoring"),
        ],
        [Inches(2.2), Inches(4.8)],
    )


def add_ai_roadmap(document):
    document.add_heading("7. AI Roadmap", level=1)
    add_table(
        document,
        ["Phase", "AI Capabilities"],
        [
            ("Phase 1", "MeghaBot; AI Notes; AI Dashboard Insights"),
            ("Phase 2", "Citizen trend forecasting; AI-based priority scoring"),
        ],
        [Inches(1.5), Inches(5.5)],
    )


def add_roles(document):
    document.add_heading("8. User Roles and Access Overview", level=1)
    add_table(
        document,
        ["Role", "Primary Access / Responsibility"],
        [
            ("Citizen", "Self-registration, login, dashboard view, appointment creation, scheme application, gate pass download, grievance raising."),
            ("DEO", "Citizen registration, walk-in appointment creation, citizen update, document upload, citizen photo capture, public identification."),
            ("CMO", "Review submitted appointments, request missing information, edit category, add remarks, forward to Approver, reports and dashboards."),
            ("Approver", "Review CMO forwarded appointments, add Jt. Sec. remarks, forward to department, approve, reject, schedule, or mark follow-up."),
            ("HCM", "Review approved or scheduled appointments, accept/modify/reject/delay using swipe-card action flow, add decisions and remarks."),
            ("OSD", "Schedule/calendar visibility, reports, HCM workflow support, event time changes where permitted."),
            ("Admin", "Full system access, staff user management, scheme management, appointment type configuration, system administration."),
        ],
        [Inches(1.3), Inches(5.7)],
    )


def add_common_requirements(document):
    document.add_heading("9. Common Functional Requirements", level=1)
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("COM-001", "The system shall enforce role-based access for all staff modules."),
            ("COM-002", "The system shall maintain audit trail information for administrative and reportable actions wherever applicable."),
            ("COM-003", "The system shall use reference tables for configurable values such as citizen designation, agenda type, scheme type, project category, and department lists."),
            ("COM-004", "The system shall generate unique application or appointment numbers after successful submission."),
            ("COM-005", "The system shall allow uploaded documents to be viewed or downloaded by authorized staff."),
            ("COM-006", "The system shall support status-based action enablement, including gate pass download only after approval/scheduling criteria are met."),
            ("COM-007", "The system shall support third-party integration fallback where identity service downtime occurs, with KYC marked as Pending."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_citizen_module(document):
    document.add_heading("10. Citizen Module", level=1)
    document.add_paragraph(
        "The Citizen Module enables citizen registration, authentication, dashboard access, appointment creation, scheme application, and gate pass download."
    )

    document.add_heading("10.1 Citizen Registration", level=2)
    document.add_paragraph(
        "New registration shall support three identity paths: EPIC as primary ID, Aadhaar as secondary ID if EPIC is not available, and No ID for quick registration or cases where the citizen does not have the identity proof at the time of registration."
    )

    document.add_heading("EPIC Registration Flow", level=3)
    add_numbered(
        document,
        [
            "Citizen selects ID Type as EPIC and enters the EPIC number.",
            "The system validates the EPIC format: first three characters must be letters and the remaining seven characters must be digits.",
            "The backend calls the third-party EPIC service API by passing citizen name and EPIC number.",
            "The EPIC service response is used to fetch name match score, assembly constituency details, and polling details.",
            "Citizen enters an active mobile number and the system sends OTP using SMS gateway for mobile validation.",
            "Successful OTP validation completes EPIC-based KYC.",
            "The system captures and stores a live citizen photo.",
            "Citizen selects citizen designation from a predefined reference table.",
            "If the citizen is outside the state, the NA checkbox may be selected. Otherwise, address data returned by EPIC shall be auto-populated where available. If address data is not returned, the user may manually select District, Constituency, and Booth/Village.",
            "Citizen selects Agenda Type from a reference table.",
            "Citizen clicks Complete Registration to finish registration.",
        ],
    )

    document.add_heading("Aadhaar Registration Flow", level=3)
    add_numbered(
        document,
        [
            "Citizen selects ID Type as Aadhaar.",
            "The system displays a QR code for Aadhaar mobile app based authentication.",
            "Citizen authenticates Aadhaar in the mobile app and scans the QR code.",
            "The system calls the OVSC third-party API after authentication.",
            "The OVSC response is used to fetch citizen photo, name, mobile number, and address.",
            "The Aadhaar photo returned by the service shall be stored; live photo capture is not required for this path.",
            "Citizen selects citizen designation from a predefined reference table.",
            "If the citizen is outside the state, the NA checkbox may be selected. Otherwise, address data returned by Aadhaar shall be auto-populated where available. If address data is not returned, the user may manually select District, Constituency, and Booth/Village.",
            "Citizen selects Agenda Type from a reference table.",
            "Citizen clicks Complete Registration to finish registration.",
        ],
    )

    document.add_heading("No ID Registration Flow", level=3)
    add_numbered(
        document,
        [
            "Citizen selects ID Type as No ID.",
            "The system validates citizen name and mobile number.",
            "The system captures and stores a live citizen photo.",
            "Citizen selects citizen designation from a predefined reference table.",
            "If the citizen is outside the state, the NA checkbox may be selected. Otherwise, the user manually selects District, Constituency, and Booth/Village.",
            "Citizen selects Agenda Type from a reference table.",
            "Citizen clicks Complete Registration to finish registration.",
        ],
    )

    document.add_heading("KYC Fallback Rule", level=3)
    add_bullets(
        document,
        [
            "If EPIC or Aadhaar third-party services are down, the system shall allow registration without completed KYC.",
            "Such registrations shall be marked as KYC Pending for later completion.",
        ],
    )

    document.add_heading("10.2 Citizen Login", level=2)
    add_numbered(
        document,
        [
            "Citizen enters mobile number.",
            "The system sends OTP through SMS gateway.",
            "If the mobile number is not registered, the system provides an option to register.",
            "If the mobile number is linked to more than one registered citizen, the system asks for EPIC number to identify the exact citizen profile.",
            "After successful authentication, the citizen is redirected to the citizen dashboard.",
        ],
    )

    document.add_heading("10.3 Citizen Dashboard", level=2)
    add_bullets(
        document,
        [
            "Dashboard shall show appointments, total visits, active schemes, and total grievances.",
            "Dashboard shall show visitor profile and personal details.",
            "Appointment history shall show submitted applications or appointments, current status, view option, and gate pass download option where enabled.",
            "Menu options shall include Create Appointment, Apply for Scheme, and Raise Grievance.",
        ],
    )

    document.add_heading("10.4 Citizen Appointment", level=2)
    add_numbered(
        document,
        [
            "Step 1 - Citizen Details: system auto-populates name, mobile number, EPIC/Aadhaar details, constituency, and KYC status.",
            "Step 2 - Agenda: system auto-populates agenda type and brief description captured during registration, with option to change.",
            "Step 3 - Scheme Details: when agenda type is Scheme, the scheme form is enabled.",
            "Scheme details include Scheme Type from reference table, Application Type, Project Category, Project Name, Who Will Benefit, Total People Benefited, Estimated Cost, Community Contribution, Justification, MLA Approval Yes/No, and optional Scheme Taken in Last 2 Years.",
            "Step 4 - Associates: user may select a checkbox to search existing registered citizens and add them as associates when a group is visiting. Every associate must be registered.",
            "Step 5 - Documents: user may optionally upload scanned scheme form or supporting documents. DEO may upload later if required.",
            "Step 6 - Review: user reviews and submits the appointment.",
            "System generates an application number after submission.",
        ],
    )

    document.add_heading("10.5 Citizen Apply for Scheme", level=2)
    add_numbered(
        document,
        [
            "Step 1 - Scheme Selection: citizen selects scheme type from reference table.",
            "Step 2 - Project Details: citizen enters Project Category, Project Name, Who Will Benefit, Total People Benefited, Estimated Cost, Community Contribution, Justification, MLA Approval Yes/No, and optional Scheme Taken in Last 2 Years.",
            "Step 3 - Financial Details: citizen adds item description, quantity, and unit cost. Multiple line items can be added.",
            "Step 4 - Community Contribution: citizen records contribution details.",
            "Step 5 - Documents: citizen optionally uploads scanned scheme form or supporting documents. DEO may upload later if required.",
            "Step 6 - Review: citizen reviews and submits the scheme application.",
            "System generates an application number after submission.",
        ],
    )

    document.add_heading("10.6 Download the Pass", level=2)
    add_bullets(
        document,
        [
            "From appointment history, Download Pass shall be enabled only when appointment status is approved or otherwise eligible for entry.",
            "Citizen can use the pass for entry when appointment is fixed.",
            "The QR code on the pass shall expire after appointment completion.",
        ],
    )

    document.add_heading("Citizen Module Requirements", level=2)
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("CIT-001", "The system shall support citizen registration through EPIC, Aadhaar, and No ID paths."),
            ("CIT-002", "The system shall validate EPIC number format before calling EPIC service."),
            ("CIT-003", "The system shall integrate with EPIC service for name score match, constituency, and polling details."),
            ("CIT-004", "The system shall integrate with SMS gateway for mobile OTP validation."),
            ("CIT-005", "The system shall integrate with Aadhaar OVSC service using QR authentication flow."),
            ("CIT-006", "The system shall store live photo for EPIC and No ID registrations and Aadhaar photo for Aadhaar registrations."),
            ("CIT-007", "The system shall support KYC Pending registration when identity services are unavailable."),
            ("CIT-008", "The system shall provide citizen dashboard with profile, appointments, schemes, visits, grievances, and appointment history."),
            ("CIT-009", "The system shall allow citizen appointment creation with agenda, scheme details, associates, documents, review, and application number generation."),
            ("CIT-010", "The system shall allow scheme application submission with financial line items and document upload."),
            ("CIT-011", "The system shall enable pass download only when status rules permit."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_deo_module(document):
    document.add_heading("11. DEO Module", level=1)
    add_bullets(
        document,
        [
            "DEO logs in with DEO credentials and role-based access.",
            "DEO can register citizens and create appointments for walk-in citizens using the same citizen registration and appointment flows.",
            "DEO can update citizen details.",
            "DEO can add supporting documents.",
            "DEO can capture citizen photo after CM visits, where required.",
        ],
    )
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("DEO-001", "The system shall allow DEO to perform citizen registration on behalf of citizens."),
            ("DEO-002", "The system shall allow DEO to create walk-in appointments."),
            ("DEO-003", "The system shall allow DEO to update citizen details and upload supporting documents."),
            ("DEO-004", "The system shall allow DEO to capture citizen photo or meeting proof after CM visit where required."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_public_identification(document):
    document.add_heading("12. Public Identification Module", level=1)
    document.add_paragraph("Accessible by staff roles such as DEO, CMO, Approver, and HCM.")
    add_bullets(
        document,
        [
            "Authorized staff can search citizens by phone number, EPIC, name, or district.",
            "Search results shall show citizen profile and personal information.",
            "Search results shall show Scheme and Meeting History, including schemes applied and appointment history with status.",
        ],
    )
    document.add_heading("Pending Item - Identify by Face", level=2)
    add_bullets(
        document,
        [
            "Face match API integration is pending.",
            "The future flow shall capture a live citizen photo and perform 1:N matching against the internal database.",
            "A separate API and GPU-enabled processing setup are required for fast 1:N match.",
            "Estimated time for development and integration is 2 weeks after required infrastructure availability.",
        ],
    )
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("PID-001", "The system shall allow authorized staff to search citizens using phone number, EPIC, name, or district."),
            ("PID-002", "The system shall show citizen profile, personal information, scheme history, and appointment history."),
            ("PID-003", "The system shall support future face-identification integration subject to API and GPU infrastructure readiness."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_staff_dashboard(document):
    document.add_heading("13. Staff Dashboard", level=1)
    add_bullets(
        document,
        [
            "After staff login, the system redirects to staff dashboard.",
            "Dashboard shall show Today's Appointments, Active Scheme Applications, CMO Reviews for CMO login, Pending Approvals for Approver login, Pending Follow-ups for HCM login, and Walk-in Today for DEO login.",
            "Recent Activity shall be visible for all logins.",
            "Quick Actions shall show Walk-in Counter, Register Visitor, and Identify Person for DEO login.",
            "Quick Actions shall show Scheme Heatmap and View Reports for Approver, CMO, HCM, and OSD roles.",
            "Today's Schedule shall be visible for Approver, CMO, HCM, and OSD roles.",
            "Appointment by Type pie chart shall be visible for Approver, CMO, HCM, and OSD roles.",
            "Scheme Applications This Month bar chart shall be visible for Approver, CMO, HCM, and OSD roles.",
            "AI Dashboard Insights shall show Top Requested Schemes, District-wise Applications, Top Project Categories, and AI Notes for Approver, CMO, HCM, and OSD roles.",
        ],
    )
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("DASH-001", "The system shall show dashboard cards and widgets based on logged-in staff role."),
            ("DASH-002", "The system shall show recent activity for all staff users."),
            ("DASH-003", "The system shall show role-specific quick actions."),
            ("DASH-004", "The system shall show schedule, appointment type chart, scheme application chart, and AI insights for authorized senior roles."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_calendar(document):
    document.add_heading("14. Calendar and Schedule Module", level=1)
    document.add_paragraph("Roles: CMO, Approver, OSD, HCM.")
    add_numbered(
        document,
        [
            "Calendar opens with current date in focus.",
            "Past dates shall be shown with cancelled or disabled indication.",
            "For current or future dates, authorized users can create an event by clicking Add Event.",
            "Event form supports Public Darbar, Walk-in, or Program with selected date and timings.",
            "Approver can assign waiting citizens to an event.",
            "System sends notification to assigned citizens and generates pass.",
            "Approver and OSD can change event timings by dragging the event in the calendar where permitted.",
        ],
    )
    document.add_paragraph("Pending items: No pending items.")
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("CAL-001", "The system shall show current date in focus and past dates as unavailable or cancelled."),
            ("CAL-002", "The system shall allow authorized users to create Public Darbar, Walk-in, or Program events."),
            ("CAL-003", "The system shall allow Approver to assign citizens awaiting appointment to an event."),
            ("CAL-004", "The system shall send notification and generate pass for assigned citizens."),
            ("CAL-005", "The system shall allow permitted roles to update event timings through calendar drag action."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_all_appointments(document):
    document.add_heading("15. All Appointment Module", level=1)
    document.add_paragraph("Roles: CMO and Approver.")
    add_bullets(
        document,
        [
            "Filters shall include name search, status filter, type filter, from date, and to date.",
            "Result grid shall support sorting.",
            "Grid columns shall include Applicant, Designation, Constituency, Agenda, Type, Location, Status, Created At, AI Notes, and Actions.",
            "AI Notes shall provide View and Refresh options.",
            "AI Notes View shall show AI generated Summary, Important Details, Missing Information, Risk Flags, Priority Score, Suggested Department, and AI Recommendation from uploaded documents.",
            "AI Notes Refresh shall regenerate AI notes.",
        ],
    )
    document.add_heading("Action View", level=2)
    add_bullets(
        document,
        [
            "Action View shall show personal information with photo.",
            "Action View shall show associate visitors when group visitors are linked.",
            "Action View shall show agenda and attached documents.",
            "Authorized staff can view, download, upload supporting documents, or capture visitor meeting proof photo.",
            "For CMO, the view shall support Request Missing Information from Citizen.",
            "For CMO, the view shall support editing appointment category and adding CMO remarks.",
            "For CMO, the Approver button shall move appointment from Submitted to Approver Review.",
            "For Approver, the view shall support Add Jt. Sec. Remarks by writing decision, forwarding to department, adding remarks, saving multiple remarks, and editing saved remarks through pencil icon.",
            "For Approver, available action buttons shall include Approve, Reject, Schedule, and Follow-up.",
            "Appointments marked Follow-up shall remain available for Approver to assign when an event is created.",
            "CMO shall see Submitted appointments in the list.",
        ],
    )
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("APT-001", "The system shall provide appointment search and filtering by name, status, type, from date, and to date."),
            ("APT-002", "The system shall show sortable appointment grid with required columns."),
            ("APT-003", "The system shall generate, view, and refresh AI notes from uploaded documents."),
            ("APT-004", "The system shall allow CMO to request missing information, edit category, add remarks, and forward to Approver Review."),
            ("APT-005", "The system shall allow Approver to add/edit remarks, forward to departments, approve, reject, schedule, or mark follow-up."),
            ("APT-006", "The system shall support document view, download, upload, and meeting proof photo capture in appointment action view."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_scheme_applications(document):
    document.add_heading("16. CM Scheme Applications", level=1)
    document.add_paragraph("Roles: CMO, Approver, OSD, HCM.")
    add_bullets(
        document,
        [
            "Module shall show scheme-wise application counts with Approved, Rejected, and Pending statistics cards.",
            "Application list shall show Applicant, Scheme, Project, Category, Estimated Cost, HCM Approved, Created At, Status, and View option.",
            "Results grid shall support sorting.",
        ],
    )
    document.add_paragraph("Pending items: No pending items.")
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("SCH-001", "The system shall show scheme-wise statistics cards for approved, rejected, and pending counts."),
            ("SCH-002", "The system shall show a sortable list of scheme applications with required columns."),
            ("SCH-003", "The system shall provide a View action for scheme application detail review."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_reports(document):
    document.add_heading("17. Reports Module", level=1)
    document.add_paragraph("Roles: CMO, Approver, OSD, HCM.")

    document.add_heading("17.1 Scheme Heatmap", level=2)
    add_bullets(
        document,
        [
            "Module shall show Total Applications, Approved, Pending, and Approval Rate counts.",
            "District-wise scheme distribution shall be shown on Meghalaya map with district markers.",
            "Selecting a district marker shall show district-wise Total Applications, Approved, Pending, and Approval Rate.",
            "District Summary shall be shown.",
            "Map shall show bubble chart by district with size proportional to application volume.",
            "Clicking bubbles shall show detailed statistics.",
        ],
    )

    document.add_heading("17.2 Analytics and Reports", level=2)
    add_bullets(
        document,
        [
            "Meetings per Day This Week shall be shown as bar chart.",
            "Approval vs Rejection Ratio shall be shown as pie chart.",
            "Scheme-wise Application Status shall be shown as bar chart.",
            "Top Constituencies by Applications shall be shown by location.",
        ],
    )

    document.add_heading("17.3 Audit Trail", level=2)
    add_bullets(
        document,
        [
            "Audit Trail shall be available for Admin.",
            "Audit Trail report shall support export.",
        ],
    )

    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("RPT-001", "The system shall show scheme heatmap statistics at state and district levels."),
            ("RPT-002", "The system shall show Meghalaya district bubble map with click-through details."),
            ("RPT-003", "The system shall provide analytics charts for meetings, approvals, scheme status, and top constituencies."),
            ("RPT-004", "The system shall provide admin audit trail and export report capability."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_hcm_actions(document):
    document.add_heading("18. HCM Actions", level=1)
    document.add_paragraph("Roles: HCM and OSD.")
    add_bullets(
        document,
        [
            "All approved and scheduled appointments shall appear in HCM dashboard.",
            "Date selection shall filter results for the selected date.",
            "Swipe cards shall support Swipe Right to Accept/Modify and Swipe Left to Reject/Delay.",
            "During accept/modify action, HCM can write decision, forward to department, and save remarks.",
            "Multiple remarks can be written for different departments.",
            "Saved remarks shall support edit option.",
        ],
    )
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("HCM-001", "The system shall show approved and scheduled appointments in HCM dashboard."),
            ("HCM-002", "The system shall filter HCM appointments by selected date."),
            ("HCM-003", "The system shall support swipe-based accept/modify/reject/delay actions."),
            ("HCM-004", "The system shall allow HCM to add decisions, forward to department, save multiple remarks, and edit saved remarks."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_admin(document):
    document.add_heading("19. Admin Module", level=1)
    document.add_paragraph("Role: Admin.")
    add_bullets(
        document,
        [
            "Admin shall have all options enabled.",
            "Admin can manage User Management system.",
            "Admin can create, edit, and delete staff users.",
            "Admin can manage Scheme Management.",
            "Admin can add new schemes when introduced.",
            "Admin can change scheme status from Active to Inactive.",
            "Admin can configure documents required for a scheme through action wrench icon.",
            "Admin can configure Appointment Type.",
            "Appointment Type configuration shall support status change, active/inactive, and edit configuration for a particular event.",
        ],
    )
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("ADM-001", "The system shall provide full administrative access to Admin role."),
            ("ADM-002", "The system shall allow Admin to create, edit, and delete staff users."),
            ("ADM-003", "The system shall allow Admin to add and activate/inactivate schemes."),
            ("ADM-004", "The system shall allow Admin to configure scheme-specific required documents."),
            ("ADM-005", "The system shall allow Admin to configure appointment types and event-related settings."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_ai_module(document):
    document.add_heading("20. AI Module", level=1)
    add_bullets(
        document,
        [
            "MeghaBot shall guide citizens during appointment creation and service discovery by providing contextual assistance inside the citizen workflow.",
            "AI Notes shall read uploaded appointment documents and generate structured review support for CMO and Approver users.",
            "AI Notes shall include Summary, Important Details, Missing Information, Risk Flags, Priority Score, Suggested Department, and AI Recommendation.",
            "AI Dashboard Insights shall generate governance insights such as top requested schemes, district-wise applications, top project categories, demand patterns, and AI-generated notes.",
            "OCR / Document Reading shall support extraction and understanding of scanned forms or uploaded supporting documents wherever document quality permits.",
            "Face Identification shall support future public identification by comparing a live captured face against citizen records, subject to required API and GPU infrastructure readiness.",
            "Future predictive governance analytics shall support trend forecasting, AI-based priority scoring, and decision support for citizen service planning.",
            "Example AI insight: Based on appointment statistics for Meghalaya CM Office, the platform may highlight significant citizen demand for CMSD / related schemes.",
        ],
    )
    add_table(
        document,
        ["Req. ID", "Requirement"],
        [
            ("AI-001", "The system shall provide MeghaBot to guide citizens in appointment creation."),
            ("AI-002", "The system shall generate AI Notes from uploaded appointment documents including Summary, Important Details, Missing Information, Risk Flags, Priority Score, Suggested Department, and AI Recommendation."),
            ("AI-003", "The system shall generate AI dashboard insights for authorized staff roles."),
            ("AI-004", "The system shall support OCR / document reading for uploaded scanned forms and supporting documents wherever document quality permits."),
            ("AI-005", "The system shall support future face identification for public identification subject to Face Recognition API and GPU infrastructure readiness."),
            ("AI-006", "The system shall support future predictive governance analytics, including citizen trend forecasting and AI-based priority scoring."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_business_rules(document):
    document.add_heading("21. Business Rules", level=1)
    add_table(
        document,
        ["Rule ID", "Business Rule"],
        [
            ("BR-001", "Public Darbar appointments shall be bulk scheduled by Approver."),
            ("BR-002", "Multiple registered visitors can be associated with one appointment."),
            ("BR-003", "Every associate visitor must complete citizen registration before being attached to an appointment."),
            ("BR-004", "QR code on the gate pass shall expire after appointment completion."),
            ("BR-005", "EPIC/Aadhaar service downtime shall not block registration; such records shall be marked as KYC Pending."),
            ("BR-006", "Gate pass download shall be enabled only for approved or scheduled appointment status as configured."),
            ("BR-007", "Reference tables shall be used for designation, agenda type, scheme type, project category, departments, and other configurable values."),
        ],
        [Inches(1.2), Inches(5.8)],
    )


def add_pending_items(document):
    document.add_heading("22. Pending Items and Dependencies", level=1)
    add_table(
        document,
        ["Item", "Dependency / Condition", "Estimated Effort", "Remarks"],
        [
            ("SMS Integration", "SMS template approval", "3 hours after template approval", "Required for OTP and citizen notifications."),
            ("WhatsApp Integration", "WhatsApp template approval", "3 hours after template approval", "Required for WhatsApp notifications where enabled."),
            ("Mobile App", "Mobile design, API integration, app store deployment", "2 weeks for design/API integration plus 1 week for app store approval/deployment", "Timeline depends on app store approval cycle."),
            ("Face Identification", "Face match API, separate GPU for fast 1:N matching", "2 weeks after infrastructure/API readiness", "Future enhancement for Public Identification module."),
        ],
        [Inches(1.5), Inches(2.2), Inches(1.8), Inches(1.5)],
    )


def add_non_functional(document):
    document.add_heading("23. Non-Functional and Control Requirements", level=1)
    add_table(
        document,
        ["Area", "Requirement"],
        [
            ("Security", "System access shall be role-based and authenticated for staff modules."),
            ("Identity Validation", "OTP validation and KYC integrations shall be used where available as per selected ID type."),
            ("Auditability", "Administrative and reportable actions shall be traceable through audit trail where applicable."),
            ("Data Handling", "Citizen personal data, photos, identity references, documents, and remarks shall be stored only for approved application workflow purposes."),
            ("Availability", "Third-party identity service downtime shall be handled through KYC Pending registration fallback."),
            ("Usability", "Citizen and staff workflows shall provide step-wise forms, review screens, status visibility, and action enablement based on role/status."),
            ("Reporting", "Reports shall provide dashboards, charts, heatmap, audit trail, and export capability as defined in scope."),
        ],
        [Inches(1.5), Inches(5.5)],
    )


def add_acceptance(document):
    document.add_heading("24. Acceptance Criteria for Code Freeze", level=1)
    add_bullets(
        document,
        [
            "All listed modules are available as per role and workflow described in this document, except explicitly listed pending items.",
            "Citizen registration supports EPIC, Aadhaar, No ID, and KYC Pending fallback paths.",
            "Appointment and scheme application flows generate application numbers after submission.",
            "Role-based staff actions are available for DEO, CMO, Approver, HCM, OSD, and Admin as described.",
            "Calendar scheduling, assignment, pass generation, reporting, and dashboard widgets follow the documented role access.",
            "Gate pass QR expiry after appointment completion is treated as mandatory business logic.",
            "Any change after approval shall follow change request and impact assessment process.",
        ],
    )


def build_document():
    document = Document()
    enable_update_fields_on_open(document)
    configure_document(document)
    add_cover_page(document)
    add_toc(document)
    add_document_control(document)
    add_intro(document)
    add_ai_transformation_vision(document)
    add_current_workflow_challenges(document)
    add_ai_architecture_overview(document)
    add_ai_benefits_for_government(document)
    add_ai_roadmap(document)
    add_roles(document)
    add_common_requirements(document)
    add_citizen_module(document)
    add_deo_module(document)
    add_public_identification(document)
    add_staff_dashboard(document)
    add_calendar(document)
    add_all_appointments(document)
    add_scheme_applications(document)
    add_reports(document)
    add_hcm_actions(document)
    add_admin(document)
    add_ai_module(document)
    add_business_rules(document)
    add_pending_items(document)
    add_non_functional(document)
    add_acceptance(document)

    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    document.add_heading("25. Final Approval", level=1)
    document.add_paragraph(
        "By signing below, the approving authority confirms that the functional scope, flows, pending items, "
        "business rules, and acceptance criteria in this document are approved for code freeze."
    )
    add_table(
        document,
        ["Approver Name", "Designation", "Signature", "Date"],
        [("", "", "", ""), ("", "", "", "")],
        [Inches(2), Inches(2), Inches(1.5), Inches(1.5)],
    )

    try:
        document.save(OUTPUT_PATH)
        return OUTPUT_PATH
    except PermissionError:
        document.save(FALLBACK_OUTPUT_PATH)
        return FALLBACK_OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
